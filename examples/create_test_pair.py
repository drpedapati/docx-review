#!/usr/bin/env python3
"""Create a pair of test .docx files for diff testing."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

out_dir = os.path.dirname(os.path.abspath(__file__))

# ── Old document ────────────────────────────────────────

doc_old = Document()

# Metadata
doc_old.core_properties.title = "Study Protocol Draft v1"
doc_old.core_properties.author = "Dr. Smith"

# Content
doc_old.add_heading("Introduction", level=1)
doc_old.add_paragraph(
    "This study examines the effects of methylphenidate on attention "
    "in subjects with ADHD. The methodology was applied to all subjects "
    "recruited from Cincinnati Children's Hospital."
)
doc_old.add_paragraph(
    "Previous research has shown mixed results regarding stimulant "
    "medication efficacy in pediatric populations."
)
doc_old.add_heading("Methods", level=1)
doc_old.add_paragraph(
    "Participants were recruited between January 2023 and December 2023. "
    "Inclusion criteria included a confirmed diagnosis of ADHD."
)
doc_old.add_paragraph(
    "This paragraph will be deleted in the new version."
)
doc_old.add_heading("Results", level=1)
p = doc_old.add_paragraph()
run = p.add_run("The primary outcome measure showed significant improvement ")
run = p.add_run("(p < 0.001)")
run.bold = True
run = p.add_run(" in the treatment group compared to placebo.")

doc_old.add_heading("Discussion", level=1)
doc_old.add_paragraph(
    "These findings support the use of methylphenidate for ADHD treatment. "
    "Further research is needed to determine optimal dosing."
)

old_path = os.path.join(out_dir, "test_old.docx")
doc_old.save(old_path)
print(f"Created: {old_path}")

# ── New document (modified version) ─────────────────────

doc_new = Document()

# Changed metadata
doc_new.core_properties.title = "Study Protocol Final Draft"
doc_new.core_properties.author = "Dr. Smith"

# Content with changes
doc_new.add_heading("Introduction", level=1)
doc_new.add_paragraph(
    "This study examines the effects of methylphenidate on attention "
    "in participants with ADHD. The methods were applied to all participants "
    "recruited from Cincinnati Children's Hospital Medical Center."
)  # Changed: "subjects" → "participants", "methodology" → "methods", added "Medical Center"

doc_new.add_paragraph(
    "Previous research has shown mixed results regarding stimulant "
    "medication efficacy in pediatric populations."
)  # Unchanged

doc_new.add_heading("Methods", level=1)
doc_new.add_paragraph(
    "Participants were recruited between January 2023 and June 2024. "
    "Inclusion criteria included a confirmed DSM-5 diagnosis of ADHD."
)  # Changed: "December 2023" → "June 2024", added "DSM-5"

# Deleted paragraph removed

doc_new.add_heading("Results", level=1)
p = doc_new.add_paragraph()
run = p.add_run("The primary outcome measure showed significant improvement ")
run = p.add_run("(p < 0.001)")
run.bold = True
run.italic = True  # Added italic formatting
run = p.add_run(" in the treatment group compared to placebo.")

# New paragraph added
doc_new.add_paragraph(
    "Secondary outcomes also demonstrated improvement in executive "
    "function measures (Table 2)."
)

doc_new.add_heading("Discussion", level=1)
doc_new.add_paragraph(
    "These findings support the use of methylphenidate for ADHD treatment "
    "in pediatric populations aged 6-17. Further research is needed to "
    "determine optimal dosing strategies and long-term outcomes."
)  # Modified

# New section added
doc_new.add_heading("Limitations", level=1)
doc_new.add_paragraph(
    "This study has several limitations including sample size "
    "and single-site recruitment."
)

new_path = os.path.join(out_dir, "test_new.docx")
doc_new.save(new_path)
print(f"Created: {new_path}")
print("Done! Run: docx-review --diff test_old.docx test_new.docx")
